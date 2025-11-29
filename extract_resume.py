from docx import Document
import sys

# Read the .docx file
doc = Document('omchaurasia.docx')

# Extract all text
resume_text = []
for para in doc.paragraphs:
    if para.text.strip():
        resume_text.append(para.text)

# Also extract from tables if any
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                resume_text.append(cell.text)

# Write to file with UTF-8 encoding
with open('resume_content.txt', 'w', encoding='utf-8') as f:
    for line in resume_text:
        f.write(line + '\n')

print("Resume content extracted successfully!")
print("\n=== RESUME CONTENT ===\n")
for line in resume_text:
    # Replace problematic characters for console output
    safe_line = line.encode('ascii', 'replace').decode('ascii')
    print(safe_line)
